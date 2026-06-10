# core/__init__.py


"""
Internal Website Crawler — Full Content Edition
================================================
Extracts text, headings, paragraphs, tables, numbers/data, images
(with OCR), downloadable files (PDF/DOCX/XLSX), code blocks, lists,
metadata, and structured data (JSON-LD / Open Graph).

Usage:
    python crawler.py --url https://your-internal-site.com
    python crawler.py --url https://your-internal-site.com \\
        --output data.json --depth 3 --delay 0.5 --images --ocr

Requirements (install all at once):
    pip install requests beautifulsoup4 lxml pillow pytesseract \\
                pdfplumber python-docx openpyxl

OCR also needs the Tesseract binary:
    macOS  : brew install tesseract
    Ubuntu : sudo apt install tesseract-ocr
    Windows: https://github.com/UB-Mannheim/tesseract/wiki
"""

import argparse
import base64
import io
import json
import logging
import re
import time
from collections import deque
from datetime import datetime
from pathlib import Path
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup

# ── Optional imports (graceful fallback if not installed) ──────────────────────
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx as python_docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

# ── Logging ────────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("crawler")

# ── Constants ──────────────────────────────────────────────────────────────────
IMAGE_EXTS   = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tiff", ".svg"}
DOC_EXTS     = {".pdf", ".docx", ".xlsx", ".xls", ".csv", ".txt", ".md"}
NUMERIC_RE   = re.compile(r"""
    (?:
        [$€£¥₹]?\s*\d[\d,]*(?:\.\d+)?(?:\s*[%KkMmBbTt])?   # money / percentages / K/M/B
      | \d{1,2}[/-]\d{1,2}[/-]\d{2,4}                        # dates
      | \d+\.\d+                                              # decimals
    )
""", re.VERBOSE)


# ── URL utilities ──────────────────────────────────────────────────────────────

def normalise(url: str) -> str:
    p = urlparse(url)
    clean = p._replace(fragment="").geturl()
    return clean.rstrip("/")


def same_domain(url: str, base: str) -> bool:
    return urlparse(url).netloc == urlparse(base).netloc


def url_extension(url: str) -> str:
    return Path(urlparse(url).path).suffix.lower()


# ── Table extraction ───────────────────────────────────────────────────────────

def extract_tables(soup: BeautifulSoup) -> list[dict]:
    """
    Returns a list of table dicts:
      { "headers": [...], "rows": [[...], ...], "caption": str, "text": str }
    """
    tables = []
    for table in soup.find_all("table"):
        caption_tag = table.find("caption")
        caption = caption_tag.get_text(strip=True) if caption_tag else ""

        headers = [
            th.get_text(strip=True)
            for th in table.find_all("th")
        ]

        rows = []
        for tr in table.find_all("tr"):
            cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
            if any(cells):
                rows.append(cells)

        # Build a readable plain-text version for embedding
        lines = []
        if caption:
            lines.append(f"Table: {caption}")
        if headers:
            lines.append(" | ".join(headers))
            lines.append("-" * (sum(len(h) for h in headers) + 3 * len(headers)))
        for row in rows:
            lines.append(" | ".join(row))

        tables.append({
            "caption": caption,
            "headers": headers,
            "rows":    rows,
            "text":    "\n".join(lines),
        })
    return tables


# ── List extraction ────────────────────────────────────────────────────────────

def extract_lists(soup: BeautifulSoup) -> list[dict]:
    result = []
    for lst in soup.find_all(["ul", "ol"]):
        items = [li.get_text(strip=True) for li in lst.find_all("li", recursive=False) if li.get_text(strip=True)]
        if items:
            result.append({
                "type":  lst.name,   # "ul" or "ol"
                "items": items,
                "text":  "\n".join(f"• {i}" if lst.name == "ul" else f"{n+1}. {i}"
                                   for n, i in enumerate(items)),
            })
    return result


# ── Code block extraction ──────────────────────────────────────────────────────

def extract_code_blocks(soup: BeautifulSoup) -> list[dict]:
    blocks = []
    for pre in soup.find_all("pre"):
        code = pre.find("code")
        lang = ""
        if code:
            classes = code.get("class", [])
            lang_classes = [c for c in classes if c.startswith("language-")]
            lang = lang_classes[0].replace("language-", "") if lang_classes else ""
            text = code.get_text()
        else:
            text = pre.get_text()
        if text.strip():
            blocks.append({"language": lang, "code": text.strip()})
    return blocks


# ── Number / data extraction ───────────────────────────────────────────────────

def extract_numbers(text: str) -> list[str]:
    return list(dict.fromkeys(NUMERIC_RE.findall(text)))  # deduplicated, order-preserving


# ── Metadata & structured data ─────────────────────────────────────────────────

def extract_meta(soup: BeautifulSoup) -> dict:
    meta = {}

    # Standard meta tags
    for tag in soup.find_all("meta"):
        name    = tag.get("name") or tag.get("property") or ""
        content = tag.get("content", "")
        if name and content:
            meta[name.lower()] = content

    # Open Graph
    og = {
        k.replace("og:", ""): v
        for k, v in meta.items()
        if k.startswith("og:")
    }
    if og:
        meta["open_graph"] = og

    # JSON-LD structured data
    json_ld_list = []
    for script in soup.find_all("script", type="application/ld+json"):
        try:
            json_ld_list.append(json.loads(script.string or "{}"))
        except json.JSONDecodeError:
            pass
    if json_ld_list:
        meta["json_ld"] = json_ld_list

    return meta


# ── Image handling ─────────────────────────────────────────────────────────────

def process_image(img_url: str, session: requests.Session, run_ocr: bool) -> dict:
    """Download an image and optionally OCR it."""
    record: dict = {
        "url":         img_url,
        "alt":         "",          # filled by caller
        "ocr_text":    "",
        "base64":      "",
        "error":       "",
    }
    ext = url_extension(img_url)
    if ext == ".svg":
        record["note"] = "SVG — skipped binary download"
        return record

    try:
        resp = session.get(img_url, timeout=15, stream=True)
        resp.raise_for_status()
        raw = resp.content

        # Embed as base64 (useful if you want to store images in JSON)
        record["base64"] = base64.b64encode(raw).decode()

        if run_ocr and OCR_AVAILABLE:
            img = Image.open(io.BytesIO(raw)).convert("RGB")
            record["ocr_text"] = pytesseract.image_to_string(img).strip()

    except Exception as exc:
        record["error"] = str(exc)

    return record


# ── File handlers ──────────────────────────────────────────────────────────────

def extract_pdf(raw: bytes) -> dict:
    if not PDF_AVAILABLE:
        return {"error": "pdfplumber not installed"}
    result: dict = {"pages": [], "text": "", "tables": []}
    try:
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                page_tables = page.extract_tables() or []
                result["pages"].append({
                    "page":   i,
                    "text":   page_text,
                    "tables": page_tables,
                })
                result["text"] += page_text + "\n"
                result["tables"].extend(page_tables)
    except Exception as exc:
        result["error"] = str(exc)
    return result


def extract_docx(raw: bytes) -> dict:
    if not DOCX_AVAILABLE:
        return {"error": "python-docx not installed"}
    result: dict = {"paragraphs": [], "tables": [], "text": ""}
    try:
        doc = python_docx.Document(io.BytesIO(raw))
        result["paragraphs"] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            result["tables"].append(rows)
        result["text"] = "\n".join(result["paragraphs"])
    except Exception as exc:
        result["error"] = str(exc)
    return result


def extract_xlsx(raw: bytes) -> dict:
    if not XLSX_AVAILABLE:
        return {"error": "openpyxl not installed"}
    result: dict = {"sheets": [], "text": ""}
    try:
        wb = openpyxl.load_workbook(io.BytesIO(raw), data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    rows.append(cells)
            result["sheets"].append({"name": sheet_name, "rows": rows})
            result["text"] += f"\n[Sheet: {sheet_name}]\n"
            result["text"] += "\n".join(" | ".join(r) for r in rows)
    except Exception as exc:
        result["error"] = str(exc)
    return result


def handle_file(url: str, session: requests.Session) -> dict | None:
    """Download and parse a non-HTML file linked from a page."""
    ext = url_extension(url)
    if ext not in DOC_EXTS:
        return None
    try:
        resp = session.get(url, timeout=30)
        resp.raise_for_status()
        raw = resp.content
    except Exception as exc:
        return {"url": url, "type": ext, "error": str(exc)}

    record: dict = {"url": url, "type": ext, "crawled_at": datetime.utcnow().isoformat() + "Z"}
    if ext == ".pdf":
        record.update(extract_pdf(raw))
    elif ext == ".docx":
        record.update(extract_docx(raw))
    elif ext in (".xlsx", ".xls"):
        record.update(extract_xlsx(raw))
    elif ext in (".csv", ".txt", ".md"):
        record["text"] = raw.decode("utf-8", errors="replace")
    return record


# ── Main page extractor ────────────────────────────────────────────────────────

def extract_page(
    url: str,
    session: requests.Session,
    process_images: bool = False,
    run_ocr: bool = False,
    download_files: bool = True,
    timeout: int = 15,
) -> dict | None:
    """
    Fetch a URL and return a rich content dict.

    Top-level keys:
      url, title, meta, headings, paragraphs, text, numbers,
      tables, lists, code_blocks, images, files, links, crawled_at
    """
    try:
        resp = session.get(url, timeout=timeout)
        resp.raise_for_status()
    except requests.RequestException as exc:
        log.warning("  SKIP %s — %s", url, exc)
        return None

    content_type = resp.headers.get("Content-Type", "")

    # ── Non-HTML files hit directly (e.g. crawl started on a PDF URL) ─────────
    if "html" not in content_type:
        ext = url_extension(url)
        if ext in DOC_EXTS:
            log.info("  FILE %s", url)
            return handle_file(url, session)
        log.info("  SKIP %s — unsupported type (%s)", url, content_type.split(";")[0])
        return None

    soup = BeautifulSoup(resp.text, "lxml")

    # ── Strip boilerplate ──────────────────────────────────────────────────────
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    # ── Tables & lists before stripping nav/footer (they may live there) ──────
    tables      = extract_tables(soup)
    lists       = extract_lists(soup)
    code_blocks = extract_code_blocks(soup)

    # Strip nav/footer AFTER capturing tables
    for tag in soup(["header", "footer", "nav", "aside"]):
        tag.decompose()

    # ── Title ─────────────────────────────────────────────────────────────────
    title = soup.title.get_text(strip=True) if soup.title else ""

    # ── Meta / structured data ─────────────────────────────────────────────────
    meta = extract_meta(soup)

    # ── Headings ──────────────────────────────────────────────────────────────
    headings = []
    for h in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
        txt = h.get_text(strip=True)
        if txt:
            headings.append({"level": h.name, "text": txt})

    # ── Paragraphs ────────────────────────────────────────────────────────────
    paragraphs = [p.get_text(strip=True) for p in soup.find_all("p") if p.get_text(strip=True)]

    # ── Full body text ────────────────────────────────────────────────────────
    body = soup.find("body")
    raw_text = body.get_text(separator="\n", strip=True) if body else ""
    lines     = [ln for ln in raw_text.splitlines() if ln.strip()]
    full_text = "\n".join(lines)

    # ── Numbers / data ────────────────────────────────────────────────────────
    numbers = extract_numbers(full_text)

    # ── Images ────────────────────────────────────────────────────────────────
    images = []
    for img in soup.find_all("img"):
        src = img.get("src") or img.get("data-src") or img.get("data-lazy-src") or ""
        if not src:
            continue
        abs_src = normalise(urljoin(url, src))
        record = {
            "url":      abs_src,
            "alt":      img.get("alt", "").strip(),
            "title":    img.get("title", "").strip(),
            "width":    img.get("width", ""),
            "height":   img.get("height", ""),
            "ocr_text": "",
            "base64":   "",
            "error":    "",
        }
        if process_images and abs_src.startswith("http"):
            fetched = process_image(abs_src, session, run_ocr)
            record.update(fetched)
        images.append(record)

    # ── Downloadable files linked on the page ──────────────────────────────────
    files = []
    if download_files:
        for a in soup.find_all("a", href=True):
            href = normalise(urljoin(url, a["href"]))
            ext  = url_extension(href)
            if ext in DOC_EXTS and href.startswith("http"):
                log.info("    ↳ downloading file: %s", href)
                fdata = handle_file(href, session)
                if fdata:
                    files.append(fdata)

    # ── Outgoing links ────────────────────────────────────────────────────────
    links = []
    for a in soup.find_all("a", href=True):
        href = normalise(urljoin(url, a["href"]))
        if href.startswith("http"):
            links.append(href)
    links = sorted(set(links))

    return {
        "url":         url,
        "title":       title,
        "meta":        meta,
        "headings":    headings,
        "paragraphs":  paragraphs,
        "text":        full_text,
        "numbers":     numbers,
        "tables":      tables,
        "lists":       lists,
        "code_blocks": code_blocks,
        "images":      images,
        "files":       files,
        "links":       links,
        "crawled_at":  datetime.utcnow().isoformat() + "Z",
    }


# ── Crawler ────────────────────────────────────────────────────────────────────

def crawl(
    start_url:      str,
    max_depth:      int   = 3,
    delay:          float = 0.3,
    max_pages:      int   = 500,
    output_file:    str   = "crawl_output.json",
    process_images: bool  = False,
    run_ocr:        bool  = False,
    download_files: bool  = True,
) -> list[dict]:
    start_url = normalise(start_url)
    visited: set[str]              = set()
    results: list[dict]            = []
    queue:   deque[tuple[str,int]] = deque([(start_url, 0)])

    session = requests.Session()
    session.headers.update({
        "User-Agent": "InternalCrawlerBot/2.0 (internal use; contact: your-team@company.com)"
    })

    log.info(
        "Starting crawl → %s  (depth=%d, max_pages=%d, images=%s, ocr=%s, files=%s)",
        start_url, max_depth, max_pages, process_images, run_ocr, download_files,
    )

    while queue and len(results) < max_pages:
        url, depth = queue.popleft()
        if url in visited:
            continue
        visited.add(url)

        log.info("[%d/%d] depth=%d  %s", len(results) + 1, max_pages, depth, url)

        page = extract_page(
            url, session,
            process_images=process_images,
            run_ocr=run_ocr,
            download_files=download_files,
        )
        if page is None:
            continue

        results.append(page)

        if depth < max_depth:
            for link in page.get("links", []):
                if link not in visited and same_domain(link, start_url):
                    queue.append((link, depth + 1))

        if len(results) % 10 == 0:
            _save(results, output_file)
            log.info("  ↳ auto-saved %d pages → %s", len(results), output_file)

        time.sleep(delay)

    _save(results, output_file)
    log.info("Done. Crawled %d pages. Output → %s", len(results), output_file)
    return results


def _save(results: list[dict], path: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)


# ── CLI ────────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Full-content internal website crawler → JSON",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument("--url",        required=True,           help="Start URL")
    parser.add_argument("--output",     default="crawl_output.json")
    parser.add_argument("--depth",      type=int,   default=3,   help="Max link depth")
    parser.add_argument("--delay",      type=float, default=0.3, help="Seconds between requests")
    parser.add_argument("--max-pages",  type=int,   default=500)
    parser.add_argument("--images",     action="store_true",     help="Download & embed images as base64")
    parser.add_argument("--ocr",        action="store_true",     help="OCR images (requires --images + tesseract)")
    parser.add_argument("--no-files",   action="store_true",     help="Skip downloading linked PDFs/DOCX/XLSX")

    args = parser.parse_args()

    if args.ocr and not OCR_AVAILABLE:
        log.warning("--ocr requested but Pillow/pytesseract not installed. OCR will be skipped.")
    if args.ocr and not args.images:
        log.warning("--ocr requires --images; enabling --images automatically.")
        args.images = True

    crawl(
        start_url=args.url,
        max_depth=args.depth,
        delay=args.delay,
        max_pages=args.max_pages,
        output_file=args.output,
        process_images=args.images,
        run_ocr=args.ocr,
        download_files=not args.no_files,
    )


if __name__ == "__main__":
    main()
